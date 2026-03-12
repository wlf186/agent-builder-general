#!/usr/bin/env python3
"""
AB-docx Skill - Word 文档处理

此脚本是 AB-docx Skill 的统一入口，支持多种 Word 文档处理操作。

Usage:
    python main.py <input_file> [--action <action>] [--output <output_file>] [--data <json_data>]

Actions:
    extract_text    - 提取 Word 文档文本内容
    create_document - 创建新的 Word 文档（需要 --data 参数）
    convert_to_pdf  - 将 Word 文档转换为 PDF
    get_info        - 获取文档基本信息

Example:
    python main.py ./input/document.docx --action extract_text
    python main.py ./output/new_doc.docx --action create_document --data '{"title": "测试文档", "content": ["段落1", "段落2"]}'
    python main.py ./input/document.docx --action convert_to_pdf --output ./output/
    python main.py ./input/document.docx --action get_info
"""

import argparse
import json
import sys
import os
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime

# 尝试导入 DOCX 处理库
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


def main():
    parser = argparse.ArgumentParser(
        description='AB-docx: Word 文档处理技能',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python main.py ./input/document.docx --action extract_text
  python main.py ./input/document.docx --action get_info
  python main.py ./output/new_doc.docx --action create_document --data '{"title": "测试"}'
  python main.py ./input/document.docx --action convert_to_pdf --output ./output/
        """
    )

    parser.add_argument('input_file', nargs='?', help='输入/输出 DOCX 文件路径')
    parser.add_argument('--action',
                        choices=['extract_text', 'create_document', 'convert_to_pdf', 'get_info'],
                        default='extract_text',
                        help='执行的操作类型 (default: extract_text)')
    parser.add_argument('--output', help='输出文件或目录路径')
    parser.add_argument('--data', help='JSON 格式的数据（用于创建文档）')
    parser.add_argument('--verbose', '-v', action='store_true', help='显示详细输出')

    args = parser.parse_args()

    # create_document 可以不需要 input_file
    if args.action != 'create_document':
        if not args.input_file:
            print(json.dumps({
                "status": "error",
                "error": "需要指定输入文件路径",
                "action": args.action
            }, ensure_ascii=False))
            sys.exit(1)

        # 验证输入文件
        input_path = Path(args.input_file)
        if not input_path.exists():
            print(json.dumps({
                "status": "error",
                "error": f"输入文件不存在: {args.input_file}",
                "action": args.action
            }, ensure_ascii=False))
            sys.exit(1)

    try:
        result = process_docx(args)
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(0 if result.get("status") == "success" else 1)
    except Exception as e:
        print(json.dumps({
            "status": "error",
            "error": str(e),
            "action": args.action
        }, ensure_ascii=False))
        sys.exit(1)


def process_docx(args) -> Dict[str, Any]:
    """根据 action 处理 Word 文档"""
    action = args.action

    if action == 'extract_text':
        return extract_text(args.input_file, args.verbose)
    elif action == 'create_document':
        if not args.data:
            return {
                "status": "error",
                "error": "create_document 操作需要 --data 参数",
                "action": action
            }
        try:
            data = json.loads(args.data)
        except json.JSONDecodeError as e:
            return {
                "status": "error",
                "error": f"无效的 JSON 数据: {e}",
                "action": action
            }
        return create_document(args.input_file, data, args.verbose)
    elif action == 'convert_to_pdf':
        return convert_to_pdf(args.input_file, args.output, args.verbose)
    elif action == 'get_info':
        return get_info(args.input_file, args.verbose)
    else:
        return {
            "status": "error",
            "error": f"未知操作: {action}",
            "action": action
        }


def extract_text(input_file: str, verbose: bool = False) -> Dict[str, Any]:
    """
    提取 Word 文档文本内容

    Args:
        input_file: DOCX 文件路径
        verbose: 是否显示详细输出

    Returns:
        包含提取文本的字典
    """
    if not HAS_PYTHON_DOCX:
        return {
            "status": "error",
            "error": "缺少必要的库：请安装 python-docx",
            "action": "extract_text"
        }

    try:
        doc = Document(input_file)

        all_text = []
        paragraph_count = 0
        table_count = len(doc.tables)

        # 提取段落文本
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                all_text.append(para.text)
                paragraph_count += 1

        # 提取表格文本
        if table_count > 0:
            all_text.append("\n=== 表格内容 ===")
            for t_idx, table in enumerate(doc.tables):
                all_text.append(f"\n表格 {t_idx + 1}:")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        all_text.append(f"  {row_text}")

        combined_text = "\n".join(all_text)

        if verbose:
            print(f"[INFO] 提取了 {paragraph_count} 个段落, {table_count} 个表格", file=sys.stderr)

        return {
            "status": "success",
            "action": "extract_text",
            "output": combined_text,
            "paragraphs": paragraph_count,
            "tables": table_count,
            "char_count": len(combined_text)
        }

    except Exception as e:
        return {
            "status": "error",
            "error": f"提取文本失败: {str(e)}",
            "action": "extract_text"
        }


def create_document(output_file: Optional[str], data: Dict[str, Any], verbose: bool = False) -> Dict[str, Any]:
    """
    创建新的 Word 文档

    Args:
        output_file: 输出文件路径
        data: 文档数据（包含 title, content 等）
        verbose: 是否显示详细输出

    Returns:
        包含创建结果的字典
    """
    if not HAS_PYTHON_DOCX:
        return {
            "status": "error",
            "error": "缺少必要的库：请安装 python-docx",
            "action": "create_document"
        }

    try:
        doc = Document()

        # 添加标题
        title = data.get('title', '新建文档')
        doc.add_heading(title, level=0)

        # 添加内容
        content = data.get('content', [])
        if isinstance(content, list):
            for item in content:
                if isinstance(item, dict):
                    # 带格式的段落
                    para_type = item.get('type', 'paragraph')
                    if para_type == 'heading':
                        doc.add_heading(item.get('text', ''), level=item.get('level', 1))
                    elif para_type == 'bullet':
                        doc.add_paragraph(item.get('text', ''), style='List Bullet')
                    else:
                        doc.add_paragraph(item.get('text', ''))
                else:
                    # 简单文本
                    doc.add_paragraph(str(item))
        elif isinstance(content, str):
            doc.add_paragraph(content)

        # 添加作者信息
        author = data.get('author')
        if author:
            doc.core_properties.author = author

        # 添加创建时间
        doc.core_properties.created = datetime.now()

        # 确定输出路径
        if not output_file:
            output_file = f"new_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # 确保输出目录存在
        output_path = Path(output_file)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # 保存文档
        doc.save(output_file)

        if verbose:
            print(f"[INFO] 创建文档: {output_file}", file=sys.stderr)

        return {
            "status": "success",
            "action": "create_document",
            "output": {
                "title": title,
                "content_items": len(content) if isinstance(content, list) else 1,
                "output_file": str(output_path.absolute())
            },
            "files": [str(output_path.absolute())]
        }

    except Exception as e:
        return {
            "status": "error",
            "error": f"创建文档失败: {str(e)}",
            "action": "create_document"
        }


def convert_to_pdf(input_file: str, output_dir: Optional[str] = None, verbose: bool = False) -> Dict[str, Any]:
    """
    将 Word 文档转换为 PDF

    Args:
        input_file: DOCX 文件路径
        output_dir: 输出目录路径
        verbose: 是否显示详细输出

    Returns:
        包含转换结果的字典
    """
    try:
        # 确定输出路径
        input_path = Path(input_file)
        if not output_dir:
            output_dir = str(input_path.parent)

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        output_file = str(output_path / f"{input_path.stem}.pdf")

        # 尝试使用 docx2pdf（仅 Windows）或 LibreOffice
        converted = False
        conversion_method = None

        # 方法1: 尝试使用 docx2pdf
        try:
            from docx2pdf import convert
            convert(input_file, output_file)
            converted = True
            conversion_method = "docx2pdf"
        except ImportError:
            pass
        except Exception as e:
            if verbose:
                print(f"[WARN] docx2pdf 转换失败: {e}", file=sys.stderr)

        # 方法2: 尝试使用 LibreOffice
        if not converted:
            try:
                import subprocess
                result = subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, input_file],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0:
                    converted = True
                    conversion_method = "LibreOffice"
            except Exception as e:
                if verbose:
                    print(f"[WARN] LibreOffice 转换失败: {e}", file=sys.stderr)

        if not converted:
            return {
                "status": "error",
                "error": "无法转换为 PDF：请安装 docx2pdf（Windows）或 LibreOffice",
                "action": "convert_to_pdf"
            }

        if verbose:
            print(f"[INFO] 使用 {conversion_method} 转换成功: {output_file}", file=sys.stderr)

        return {
            "status": "success",
            "action": "convert_to_pdf",
            "output": {
                "conversion_method": conversion_method,
                "output_file": output_file
            },
            "files": [output_file]
        }

    except Exception as e:
        return {
            "status": "error",
            "error": f"转换为 PDF 失败: {str(e)}",
            "action": "convert_to_pdf"
        }


def get_info(input_file: str, verbose: bool = False) -> Dict[str, Any]:
    """
    获取 Word 文档基本信息

    Args:
        input_file: DOCX 文件路径
        verbose: 是否显示详细输出

    Returns:
        包含文档信息的字典
    """
    if not HAS_PYTHON_DOCX:
        return {
            "status": "error",
            "error": "缺少必要的库：请安装 python-docx",
            "action": "get_info"
        }

    try:
        doc = Document(input_file)
        props = doc.core_properties

        # 统计内容
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        char_count = sum(len(p.text) for p in doc.paragraphs)

        info = {
            "file_path": str(Path(input_file).absolute()),
            "file_size": Path(input_file).stat().st_size,
            "paragraphs": paragraph_count,
            "tables": table_count,
            "word_count": word_count,
            "char_count": char_count,
            "properties": {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or ""
            }
        }

        if verbose:
            print(f"[INFO] 文档信息: {paragraph_count} 段落, {table_count} 表格", file=sys.stderr)

        return {
            "status": "success",
            "action": "get_info",
            "output": info
        }

    except Exception as e:
        return {
            "status": "error",
            "error": f"获取文档信息失败: {str(e)}",
            "action": "get_info"
        }


if __name__ == "__main__":
    main()
