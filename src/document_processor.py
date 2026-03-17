"""
文档处理器
支持 PDF/DOCX/TXT/MD 文档解析与分块
"""
import logging
from pathlib import Path
from typing import List, Optional

from src.models import Chunk, DocumentStatus

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器

    负责：
    - 文档格式解析 (PDF/DOCX/TXT/MD)
    - 文本分块（使用 LangChain RecursiveCharacterTextSplitter）
    - 分块元数据生成
    """

    SUPPORTED_FORMATS = {".pdf", ".txt", ".md", ".docx"}

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ):
        """初始化文档处理器

        Args:
            chunk_size: 每块字符数
            chunk_overlap: 块之间重叠字符数
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._splitter = None  # 延迟加载

    def _get_splitter(self):
        """获取文本分块器"""
        if self._splitter is None:
            try:
                from langchain_text_splitters import RecursiveCharacterTextSplitter

                self._splitter = RecursiveCharacterTextSplitter(
                    chunk_size=self.chunk_size,
                    chunk_overlap=self.chunk_overlap,
                    separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
                    length_function=len,
                )
            except ImportError as e:
                raise ImportError(
                    "langchain_text_splitters 未安装。"
                    "请运行: pip install langchain-text-splitters"
                ) from e

        return self._splitter

    def parse(self, file_path: Path) -> str:
        """解析文档提取文本

        Args:
            file_path: 文档路径

        Returns:
            str: 提取的文本内容

        Raises:
            ValueError: 不支持的文件格式
        """
        suffix = file_path.suffix.lower()

        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"不支持的文件格式: {suffix}。"
                f"支持的格式: {', '.join(self.SUPPORTED_FORMATS)}"
            )

        logger.info(f"正在解析文档: {file_path.name}")

        try:
            if suffix == ".pdf":
                return self._parse_pdf(file_path)
            elif suffix == ".docx":
                return self._parse_docx(file_path)
            elif suffix in {".txt", ".md"}:
                return self._parse_text(file_path)
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            raise

        return ""

    def _parse_pdf(self, file_path: Path) -> str:
        """解析 PDF 文档"""
        try:
            import pypdfium2

            pdf = pypdfium2.PdfDocument(file_path)
            text_parts = []

            for page_num, page in enumerate(pdf):
                try:
                    text_page = page.get_textpage()
                    text = text_page.get_text_range()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"PDF 第 {page_num + 1} 页解析失败: {e}")

            pdf.close()

            full_text = "\n".join(text_parts)
            logger.info(f"PDF 解析完成，共 {len(text_parts)} 页，{len(full_text)} 字符")
            return full_text

        except ImportError as e:
            raise ImportError(
                "pypdfium2 未安装。请运行: pip install pypdfium2"
            ) from e
        except Exception as e:
            logger.error(f"PDF 解析错误: {e}")
            raise

    def _parse_docx(self, file_path: Path) -> str:
        """解析 DOCX 文档"""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            full_text = "\n".join(paragraphs)
            logger.info(f"DOCX 解析完成，共 {len(paragraphs)} 段，{len(full_text)} 字符")
            return full_text

        except ImportError as e:
            raise ImportError(
                "python-docx 未安装。请运行: pip install python-docx"
            ) from e
        except Exception as e:
            logger.error(f"DOCX 解析错误: {e}")
            raise

    def _parse_text(self, file_path: Path) -> str:
        """解析纯文本文件"""
        try:
            text = file_path.read_text(encoding="utf-8")
            logger.info(f"文本文件解析完成，{len(text)} 字符")
            return text
        except UnicodeDecodeError:
            # 尝试其他编码
            for encoding in ["gbk", "gb2312", "latin1"]:
                try:
                    text = file_path.read_text(encoding=encoding)
                    logger.info(f"文本文件解析完成（编码: {encoding}），{len(text)} 字符")
                    return text
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"无法解码文件: {file_path}")

    def chunk(self, text: str, doc_id: str) -> List[Chunk]:
        """文本分块

        Args:
            text: 待分块文本
            doc_id: 文档 ID

        Returns:
            List[Chunk]: 分块列表
        """
        if not text:
            return []

        splitter = self._get_splitter()
        chunks_text = splitter.split_text(text)

        chunks = []
        position = 0

        for i, chunk_text in enumerate(chunks_text):
            chunk = Chunk(
                chunk_id=f"{doc_id}_{i}",
                doc_id=doc_id,
                content=chunk_text,
                chunk_index=i,
                start_pos=position,
                end_pos=position + len(chunk_text)
            )
            chunks.append(chunk)
            position += len(chunk_text)

        logger.info(f"文本分块完成，共 {len(chunks)} 块")
        return chunks

    def process(self, file_path: Path, doc_id: str) -> tuple[str, List[Chunk]]:
        """完整处理流程：解析 → 分块

        Args:
            file_path: 文档路径
            doc_id: 文档 ID

        Returns:
            tuple[str, List[Chunk]]: (完整文本, 分块列表)
        """
        text = self.parse(file_path)
        chunks = self.chunk(text, doc_id)

        return text, chunks

    def validate_file(self, file_path: Path, max_size_mb: int = 10) -> tuple[bool, Optional[str]]:
        """验证文件是否可处理

        Args:
            file_path: 文件路径
            max_size_mb: 最大文件大小（MB）

        Returns:
            tuple[bool, Optional[str]]: (是否有效, 错误信息)
        """
        # 检查文件存在
        if not file_path.exists():
            return False, "文件不存在"

        # 检查文件格式
        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            return False, f"不支持的文件格式: {suffix}"

        # 检查文件大小
        file_size = file_path.stat().st_size
        max_size_bytes = max_size_mb * 1024 * 1024
        if file_size > max_size_bytes:
            return False, f"文件过大 ({file_size / 1024 / 1024:.1f}MB)，最大支持 {max_size_mb}MB"

        return True, None
